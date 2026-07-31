"""
数字工具箱 - 文档转换后端 (Python + LibreOffice)
==================================================
独立部署于 VPS，为 Cloudflare Pages 前端提供文档转换 API。
支持 PDF / Word / Excel 三种格式相互转换。

转换引擎：
  - LibreOffice headless (soffice)  : doc/docx↔pdf, xlsx/xls→pdf, pdf→docx
  - python-docx + openpyxl + pdfplumber : word↔excel, pdf→excel（表格提取）

技术栈：FastAPI + LibreOffice + python-docx + openpyxl + pdfplumber

启动方式：
    pip install -r requirements.txt
    uvicorn app:app --host 0.0.0.0 --port 8000

接口一览（POST，均上传字段名 file）：
    /api/doc/word2pdf      Word(.doc/.docx) → PDF
    /api/doc/pdf2word      PDF → Word(.docx)          [文本型PDF最佳]
    /api/doc/xlsx2pdf      Excel(.xls/.xlsx) → PDF
    /api/doc/word2excel    Word(.docx) → Excel(.xlsx)  [提取表格/段落]
    /api/doc/excel2word    Excel(.xls/.xlsx) → Word(.docx) [表格还原]
    /api/doc/pdf2excel     PDF → Excel(.xlsx)          [提取表格/文本]
    GET  /health           健康检查

前端对接（Vue 3 fetch 示例）：
    const fd = new FormData()
    fd.append('file', file)
    const res = await fetch(`${API_BASE}/api/doc/word2pdf`, { method: 'POST', body: fd })
    const blob = await res.blob()
"""
import asyncio
import os
import shutil
import subprocess
import uuid
from pathlib import Path

from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse

# ---------- 配置 ----------
MAX_FILE_SIZE = 50 * 1024 * 1024        # 50MB 上限
CONVERT_TIMEOUT = 180                   # LibreOffice 转换超时（秒）
ALLOWED_WORD_EXT = {".doc", ".docx"}
ALLOWED_PDF_EXT = {".pdf"}
ALLOWED_XLSX_EXT = {".xls", ".xlsx"}
# 临时目录：环境变量 TOOLBOX_TEMP_DIR 可覆盖（默认 /var/lib/toolbox-backend/tmp）
# 配套定时清理脚本 backend/cleanup_uploads.py（crontab 每 10 分钟清理超 15 分钟的文件）
TEMP_ROOT = Path(os.environ.get("TOOLBOX_TEMP_DIR", "/var/lib/toolbox-backend/tmp"))

# ---------- 应用 ----------
app = FastAPI(title="工具箱文档转换服务", version="2.0.0")

# CORS：允许前端（CF Pages 域名）跨域调用
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # 生产环境建议收敛为具体域名
    allow_methods=["POST", "GET", "OPTIONS"],
    allow_headers=["*"],
)


# ============================================================
#  Python 转换器（word↔excel、pdf→excel）
# ============================================================
import re

# XML 非法控制字符（openpyxl 写入会抛 IllegalCharacterError）
_ILLEGAL_XML_CHARS = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")


def _clean_cell(v) -> str:
    """清洗单元格值：去 None / 控制字符 / 首尾空白"""
    if v is None:
        return ""
    return _ILLEGAL_XML_CHARS.sub("", str(v)).strip()


def word_to_excel(src: Path, dst: Path) -> None:
    """Word(.docx) → Excel(.xlsx)：有表格则逐表提取，无表格则段落文本"""
    from docx import Document
    from openpyxl import Workbook

    doc = Document(str(src))
    wb = Workbook()

    if doc.tables:
        for i, table in enumerate(doc.tables):
            ws = wb.active if i == 0 else wb.create_sheet()
            ws.title = f"表格{i + 1}"[:31]
            for row in table.rows:
                ws.append([_clean_cell(cell.text) for cell in row.cells])
    else:
        ws = wb.active
        ws.title = "内容"
        for p in doc.paragraphs:
            if p.text.strip():
                ws.append([_clean_cell(p.text)])

    wb.save(str(dst))


def excel_to_word(src: Path, dst: Path) -> None:
    """Excel(.xls/.xlsx) → Word(.docx)：每个 sheet 一个标题 + 完整表格"""
    from docx import Document
    from openpyxl import load_workbook

    wb = load_workbook(str(src), data_only=True)
    doc = Document()

    for ws in wb.worksheets:
        doc.add_heading(ws.title or "Sheet", level=1)
        rows = [list(r) for r in ws.iter_rows(values_only=True)]
        rows = [r for r in rows if any(v is not None for v in r)]
        if rows:
            cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = "Table Grid"
            for i, row in enumerate(rows):
                for j in range(cols):
                    val = row[j] if j < len(row) else None
                    table.rows[i].cells[j].text = _clean_cell(val)
        doc.add_paragraph()

    doc.save(str(dst))


def pdf_to_excel(src: Path, dst: Path) -> None:
    """PDF → Excel(.xlsx)：提取页面表格（pdfplumber），无表格退化为文本行"""
    import pdfplumber
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "提取结果"
    wrote = 0

    with pdfplumber.open(str(src)) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            if tables:
                for t in tables:
                    for row in t:
                        ws.append([_clean_cell(c) for c in row])
                    ws.append([])  # 空行分隔不同表格
                    wrote += len(t)
            else:
                text = page.extract_text() or ""
                for line in text.split("\n"):
                    if line.strip():
                        ws.append([_clean_cell(line)])
                        wrote += 1

    if wrote == 0:
        raise HTTPException(status_code=422, detail="未能从 PDF 中提取到任何文字或表格（可能是扫描件）")

    wb.save(str(dst))


# ============================================================
#  基础工具函数
# ============================================================
def _ensure_libreoffice() -> None:
    """检查 soffice 是否可用，缺失时给出安装指引"""
    if shutil.which("soffice") is None and shutil.which("libreoffice") is None:
        raise HTTPException(
            status_code=500,
            detail="LibreOffice 未安装。请执行："
            "apt-get update && apt-get install -y libreoffice-writer libreoffice-calc",
        )


async def _run_soffice(src: Path, out_dir: Path, target_format: str) -> Path:
    """调用 LibreOffice headless 转换，返回输出文件路径"""
    _ensure_libreoffice()
    # 独立 HOME，避免并行任务锁冲突
    home_dir = out_dir / "home"
    home_dir.mkdir(parents=True, exist_ok=True)
    env = {**os.environ, "HOME": str(home_dir)}

    cmd = [
        "soffice", "--headless", "--norestore", "--nologo",
        "--convert-to", target_format,
        "--outdir", str(out_dir),
        str(src),
    ]
    proc = await asyncio.create_subprocess_exec(
        *cmd, stdout=asyncio.subprocess.PIPE, stderr=asyncio.subprocess.PIPE, env=env,
    )
    try:
        stdout, stderr = await asyncio.wait_for(proc.communicate(), timeout=CONVERT_TIMEOUT)
    except asyncio.TimeoutError:
        proc.kill()
        raise HTTPException(status_code=504, detail="转换超时，文件可能过大或格式复杂")

    if proc.returncode != 0:
        raise HTTPException(
            status_code=500,
            detail=f"LibreOffice 转换失败: {stderr.decode('utf-8', errors='ignore')[:500]}",
        )

    out_file = out_dir / f"{src.stem}.{target_format}"
    if not out_file.exists():
        candidates = list(out_dir.glob(f"{src.stem}.*"))
        if not candidates:
            raise HTTPException(status_code=500, detail="转换完成但未找到输出文件")
        out_file = candidates[0]
    return out_file


async def _handle_soffice(upload: UploadFile, allowed_ext: set, target_format: str, content_type: str) -> FileResponse:
    """LibreOffice 引擎转换：校验 → 保存 → 转换 → 返回"""
    await _validate(upload, allowed_ext)
    work_dir = _new_task_dir()
    try:
        src = work_dir / f"input{Path(upload.filename).suffix.lower()}"
        _save_upload(upload, src)
        out_file = await _run_soffice(src, work_dir, target_format)
        return _response(out_file, content_type, upload.filename, target_format)
    finally:
        _schedule_cleanup(work_dir)


async def _handle_python(upload: UploadFile, allowed_ext: set, target_ext: str, content_type: str, converter) -> FileResponse:
    """Python 引擎转换：校验 → 保存 → converter(src,dst) → 返回"""
    await _validate(upload, allowed_ext)
    work_dir = _new_task_dir()
    try:
        src = work_dir / f"input{Path(upload.filename).suffix.lower()}"
        _save_upload(upload, src)
        out_file = work_dir / f"output.{target_ext}"
        # 同步转换器放到线程池，避免阻塞事件循环
        try:
            await asyncio.to_thread(converter, src, out_file)
        except HTTPException:
            raise
        except Exception as e:
            # 文件损坏 / 格式异常等 → 返回友好 422 而非 500
            raise HTTPException(status_code=422, detail=f"文件解析失败，请检查文件是否完整或格式正确（{type(e).__name__}）")
        if not out_file.exists():
            raise HTTPException(status_code=500, detail="转换失败：未生成输出文件")
        return _response(out_file, content_type, upload.filename, target_ext)
    finally:
        _schedule_cleanup(work_dir)


async def _validate(upload: UploadFile, allowed_ext: set) -> None:
    if upload.size and upload.size > MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail=f"文件超过 {MAX_FILE_SIZE // 1024 // 1024}MB 限制")
    ext = Path(upload.filename or "").suffix.lower()
    if ext not in allowed_ext:
        raise HTTPException(status_code=400, detail=f"仅支持 {', '.join(sorted(allowed_ext))} 格式")


def _new_task_dir() -> Path:
    task_id = uuid.uuid4().hex[:12]
    work_dir = TEMP_ROOT / task_id
    work_dir.mkdir(parents=True, exist_ok=True)
    return work_dir


def _save_upload(upload: UploadFile, dst: Path) -> None:
    with open(dst, "wb") as f:
        shutil.copyfileobj(upload.file, f, length=1024 * 1024)


def _response(out_file: Path, content_type: str, original_name: str, target_ext: str) -> FileResponse:
    stem = Path(original_name).stem
    return FileResponse(
        path=str(out_file),
        media_type=content_type,
        filename=f"{stem}_converted.{target_ext}",
    )


def _schedule_cleanup(work_dir: Path) -> None:
    """延迟 60s 清理任务目录（确保文件流发送完成）"""

    async def _cleanup():
        await asyncio.sleep(60)
        shutil.rmtree(work_dir, ignore_errors=True)

    asyncio.create_task(_cleanup())


# ============================================================
#  路由
# ============================================================
@app.get("/health")
@app.get("/api/health")
async def health():
    """健康检查（供监控/前端探测；/api/health 为 Pages Functions 代理同域路径）"""
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    py_libs = {}
    for name, mod in [("python-docx", "docx"), ("openpyxl", "openpyxl"), ("pdfplumber", "pdfplumber")]:
        try:
            __import__(mod)
            py_libs[name] = True
        except ImportError:
            py_libs[name] = False
    return {
        "status": "ok",
        "libreoffice": bool(soffice),
        "python_libs": py_libs,
    }


@app.post("/api/doc/word2pdf")
async def word2pdf(file: UploadFile = File(...)):
    """Word → PDF"""
    return await _handle_soffice(file, ALLOWED_WORD_EXT, "pdf", "application/pdf")


@app.post("/api/doc/pdf2word")
async def pdf2word(file: UploadFile = File(...)):
    """PDF → Word（文本型 PDF 最佳，扫描件效果有限）"""
    return await _handle_soffice(file, ALLOWED_PDF_EXT, "docx",
                                 "application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.post("/api/doc/xlsx2pdf")
async def xlsx2pdf(file: UploadFile = File(...)):
    """Excel → PDF"""
    return await _handle_soffice(file, ALLOWED_XLSX_EXT, "pdf", "application/pdf")


@app.post("/api/doc/word2excel")
async def word2excel(file: UploadFile = File(...)):
    """Word → Excel（提取表格/段落）"""
    return await _handle_python(file, ALLOWED_WORD_EXT, "xlsx",
                                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                word_to_excel)


@app.post("/api/doc/excel2word")
async def excel2word(file: UploadFile = File(...)):
    """Excel → Word（表格还原为 Word 表格）"""
    return await _handle_python(file, ALLOWED_XLSX_EXT, "docx",
                                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                excel_to_word)


@app.post("/api/doc/pdf2excel")
async def pdf2excel(file: UploadFile = File(...)):
    """PDF → Excel（提取表格，退化提取文本）"""
    return await _handle_python(file, ALLOWED_PDF_EXT, "xlsx",
                                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                pdf_to_excel)


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
